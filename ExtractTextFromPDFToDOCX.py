# Description: 
# This program extracts text from a PDF file and dumps it into a newly created DOCX file.
# It is done character per character and maintains bold and/or italic formatting.
#
# How to use:
# Provide the absolute path of the PDF file to the input box, then click "Extract".
# The new DOCX file will be moved the same directory where the PDF file is.
#
# Author:
# Lucca Gonçalves

from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar
from docx import Document
from docx.shared import RGBColor
from PySimpleGUI import PySimpleGUI as sg
import os.path
import os
import sys

# The code below extracts text from a PDF file and writes it to a DOCX file
# It is done character by character and it mantains bold and/or italic formatting
def extract_text_to_document(document, complete_path):
    pg_counter = 0
    for page_layout in extract_pages(complete_path):
        pg_counter += 1
        prg = document.add_paragraph('')
        run = prg.add_run('Page ' + str(pg_counter))
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0)
        font.hidden = True
        for element in page_layout:
            if isinstance(element, LTTextContainer):
                prg = document.add_paragraph('')
                first_line = True
                for text_line in element:
                    line_beggining = True
                    for character in text_line:
                        if isinstance(character, LTChar):
                            #print(character.get_text())
                            run = prg.add_run(character.get_text())
                            if 'bold' in character.fontname:
                                run.bold = True
                            if 'Bold' in character.fontname:
                                run.bold = True
                            if 'italic' in character.fontname:
                                run.italic = True
                            if 'Italic' in character.fontname:
                                run.italic = True
                            if line_beggining == True and first_line == False:
                                if (len(last_run.text) > 1):
                                    last_run_text = last_run.text[-1]
                                else:
                                    last_run_text = last_run.text
                                if 97 <= ord(last_run_text) <= 122 and (65 <= ord(run.text[0]) <= 90 or 48 <= ord(run.text[0]) <= 57):
                                    last_run.add_break()
                            last_run = run
                            line_beggining = False
                    first_line = False
                        
    return

# Setting the layout of the GUI
sg.theme('Reddit')
layout = [
    [sg.Text('Absolute path to PDF file:'), sg.Input(key='path')],
    [sg.Button('Extract')]
]

# Generating the GUI
window = sg.Window('Extract text from PDF to DOCX', layout)

# Reading the events that happen in the GUI
while True:

    events, values = window.read()

    # In case user closes windows
    if events == sg.WINDOW_CLOSED:
        break

    # Start text extraction after user has given the input and clicked "Extract"
    if events == 'Extract':

        abs_path = os.path.abspath(values['path'].replace('"', ''))
        abs_path_dir, pdf_file = os.path.split(abs_path)

        document = Document()
        extract_text_to_document(document, abs_path)
        os.chdir(abs_path_dir)
        document.save(pdf_file.replace('pdf', 'docx'))
        break